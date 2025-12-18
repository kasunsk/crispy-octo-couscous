"""Document processing service for extracting text from various file formats."""
import os
from typing import Dict, List, Optional
from pathlib import Path
import pdfplumber
import PyPDF2
import pandas as pd
from docx import Document as DocxDocument
import logging

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Service for processing various document formats."""
    
    SUPPORTED_FORMATS = {
        'pdf': ['pdf'],
        'excel': ['xlsx', 'xls'],
        'word': ['docx', 'doc']
    }
    
    def __init__(self, upload_dir: str = "uploads"):
        """Initialize document processor."""
        self.upload_dir = Path(upload_dir)
        self.upload_dir.mkdir(exist_ok=True)
    
    def is_supported(self, filename: str) -> bool:
        """Check if file format is supported."""
        ext = filename.split('.')[-1].lower()
        return any(ext in formats for formats in self.SUPPORTED_FORMATS.values())
    
    def extract_text(self, file_path: str, file_type: str) -> str:
        """
        Extract text from document based on file type.
        
        Args:
            file_path: Path to the file
            file_type: Type of file (pdf, xlsx, docx, etc.)
        
        Returns:
            Extracted text content
        """
        file_path = Path(file_path)
        ext = file_type.lower()
        
        try:
            if ext == 'pdf':
                return self._extract_from_pdf(file_path)
            elif ext in ['xlsx', 'xls']:
                return self._extract_from_excel(file_path)
            elif ext in ['docx', 'doc']:
                return self._extract_from_word(file_path)
            else:
                raise ValueError(f"Unsupported file type: {file_type}")
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {str(e)}")
            raise
    
    def _extract_from_pdf(self, file_path: Path) -> str:
        """Extract text from PDF file."""
        text_content = []
        
        # Try pdfplumber first (better for complex PDFs)
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        text_content.append(text)
        except Exception as e:
            logger.warning(f"pdfplumber failed, trying PyPDF2: {str(e)}")
            # Fallback to PyPDF2
            try:
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    for page in pdf_reader.pages:
                        text = page.extract_text()
                        if text:
                            text_content.append(text)
            except Exception as e2:
                logger.error(f"PyPDF2 also failed: {str(e2)}")
                raise Exception(f"Failed to extract text from PDF: {str(e2)}")
        
        return "\n\n".join(text_content)
    
    def _extract_from_excel(self, file_path: Path) -> str:
        """Extract text from Excel file."""
        text_content = []
        
        try:
            # Read all sheets
            excel_file = pd.ExcelFile(file_path)
            
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(excel_file, sheet_name=sheet_name)
                
                # Add sheet name as header
                text_content.append(f"Sheet: {sheet_name}\n")
                
                # Convert DataFrame to text representation
                # Include headers and data
                text_content.append(df.to_string(index=False))
                text_content.append("\n")
        
        except Exception as e:
            logger.error(f"Error reading Excel file: {str(e)}")
            raise Exception(f"Failed to extract text from Excel: {str(e)}")
        
        return "\n\n".join(text_content)
    
    def _extract_from_word(self, file_path: Path) -> str:
        """Extract text from Word document."""
        try:
            doc = DocxDocument(file_path)
            paragraphs = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
            
            # Also extract text from tables
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        table_text.append(row_text)
                if table_text:
                    paragraphs.append("\n".join(table_text))
            
            return "\n\n".join(paragraphs)
        
        except Exception as e:
            logger.error(f"Error reading Word file: {str(e)}")
            raise Exception(f"Failed to extract text from Word document: {str(e)}")
    
    def get_file_info(self, file_path: str) -> Dict:
        """Get file information."""
        file_path = Path(file_path)
        stat = file_path.stat()
        
        return {
            "filename": file_path.name,
            "file_size": stat.st_size,
            "file_type": file_path.suffix[1:].lower() if file_path.suffix else "unknown"
        }

